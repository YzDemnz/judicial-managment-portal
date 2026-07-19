import ast
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LEGAL_TS = ROOT / "src" / "legalTerms.ts"
OUT_DOCX = ROOT / "public" / "docs" / "Judicial-Managment-Terminos-y-Condiciones.docx"
OUT_JSON = ROOT / "public" / "docs" / "legal-terms-source.json"


def parse_array_literal(source: str, name: str) -> str:
    marker = f"export const {name}"
    start = source.index(marker)
    equals_start = source.index("=", start)
    bracket_start = source.index("[", equals_start)
    depth = 0
    in_string = False
    escape = False
    quote = ""

    for index in range(bracket_start, len(source)):
      char = source[index]
      if in_string:
          if escape:
              escape = False
          elif char == "\\":
              escape = True
          elif char == quote:
              in_string = False
          continue

      if char in ("'", '"'):
          in_string = True
          quote = char
      elif char == "[":
          depth += 1
      elif char == "]":
          depth -= 1
          if depth == 0:
              return source[bracket_start:index + 1]

    raise ValueError(f"No se pudo encontrar el arreglo {name}")


def js_like_to_python(value: str):
    normalized = re.sub(r"(\n\s*)(title|body):", r"\1'\2':", value)
    normalized = re.sub(r",(\s*[\]}])", r"\1", normalized)
    return ast.literal_eval(normalized)


def read_terms():
    source = LEGAL_TS.read_text(encoding="utf-8")
    effective_match = re.search(r"legalEffectiveDate = '([^']+)'", source)
    effective_date = effective_match.group(1) if effective_match else "27 de mayo de 2026"
    intro = js_like_to_python(parse_array_literal(source, "legalIntro"))
    sections = js_like_to_python(parse_array_literal(source, "legalTermSections"))
    return effective_date, intro, sections


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(31, 41, 55)
    run.bold = bold


def style_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(17, 24, 39)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_footer(section, label: str):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run(label)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(107, 114, 128)


def add_cover(doc: Document, effective_date: str):
    for _ in range(5):
        doc.add_paragraph()

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    kicker_run = kicker.add_run("JUDICIAL MANAGMENT | MR LEGAL")
    kicker_run.bold = True
    kicker_run.font.size = Pt(11)
    kicker_run.font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("Términos y Condiciones")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    subtitle_run = subtitle.add_run("Contrato de uso para distribución controlada")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(75, 85, 99)

    meta = doc.add_table(rows=4, cols=2)
    meta.allow_autofit = False
    meta.columns[0].width = Inches(1.8)
    meta.columns[1].width = Inches(4.4)
    rows = [
        ("Producto", "Judicial Managment"),
        ("Marca", "MR Legal / Judicial Managment"),
        ("Versión del documento", "Distribución controlada 3.3.5"),
        ("Vigencia", effective_date),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_shading(row.cells[0], "E8EEF5")
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    doc.add_paragraph()
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(18)
    notice_run = notice.add_run(
        "Borrador contractual reforzado. Debe ser validado por un abogado mexicano antes de una distribución comercial abierta."
    )
    notice_run.italic = True
    notice_run.font.size = Pt(10)
    notice_run.font.color.rgb = RGBColor(107, 114, 128)

    doc.add_page_break()


def add_intro(doc: Document, intro: list[str]):
    doc.add_heading("Resumen de aceptación", level=1)
    for paragraph in intro:
        p = doc.add_paragraph(paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "FFF8E1")
    set_cell_text(
        cell,
        "Aviso: la aplicación es una herramienta de apoyo. El usuario debe conservar respaldos independientes, proteger datos de terceros y verificar toda información jurídica y salida automatizada.",
        bold=True,
    )
    doc.add_page_break()


def add_sections(doc: Document, sections):
    doc.add_heading("Contrato completo", level=1)
    for index, section in enumerate(sections):
        heading = doc.add_heading(section["title"], level=2)
        heading.paragraph_format.keep_with_next = True
        for paragraph in section["body"]:
            p = doc.add_paragraph(paragraph)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if index < len(sections) - 1:
            doc.add_page_break()


def main():
    effective_date, intro, sections = read_terms()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(
        json.dumps(
            {"effectiveDate": effective_date, "intro": intro, "sections": sections},
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    doc = Document()
    style_document(doc)
    add_footer(doc.sections[0], "Judicial Managment - Términos y Condiciones")
    add_cover(doc, effective_date)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.footer.is_linked_to_previous = False
    add_footer(body_section, "Judicial Managment - Términos y Condiciones")
    add_intro(doc, intro)
    add_sections(doc, sections)

    doc.core_properties.title = "Términos y Condiciones de Judicial Managment"
    doc.core_properties.subject = "Contrato de uso para distribución controlada"
    doc.core_properties.author = "Judicial Managment"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
